import sys
import os
import markdown
from docx import Document
from bs4 import BeautifulSoup, Tag

def convert_md_to_docx(md_path, docx_path, template_path=None):
    print(f"📄 [转换] 正在读取 Markdown 文件: {os.path.basename(md_path)}")
    
    if not os.path.exists(md_path):
        print(f"❌ [错误] 找不到 Markdown 文件: {md_path}")
        return False

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # 增加扩展：支持表格、代码块和更规范的列表解析
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'sane_lists'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # 加载模板
    if template_path and os.path.exists(template_path):
        print(f"🎨 [转换] 正在套用精装模板: {os.path.basename(template_path)}")
        try:
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
        except Exception as e:
            print(f"⚠️ [警告] 模板加载失败 ({e})，将建一个默认毛坯房...")
            doc = Document()
    else:
        print("⚠️ [转换] 根目录没找到模板，只能建个默认毛坯房了...")
        doc = Document()
    
    # 🌟 核心魔法：安全添加段落，彻底消灭软回车！
    def add_safe_paragraph(text, prefix=""):
        # 按换行符强制切开，绝不给软回车留活路
        lines = text.split('\n')
        for i, line in enumerate(lines):
            clean_line = line.strip()
            if clean_line:
                if i == 0:
                    # 第一行加上项目符号（如果有的话）
                    doc.add_paragraph(prefix + clean_line)
                else:
                    # 列表内部的换行，前面加两个空格保持视觉对齐，且作为独立段落
                    indent = "  " if prefix else ""
                    doc.add_paragraph(indent + clean_line)

    # 解析并写入 Word
    for element in soup.contents:
        if isinstance(element, Tag):
            # 获取文本时，把潜在的 <br> 标签也安全转换为 \n
            text_content = element.get_text(separator='\n').strip()
            
            if not text_content:
                continue

            if element.name == 'h1':
                doc.add_heading(text_content, level=1)
            elif element.name == 'h2':
                doc.add_heading(text_content, level=2)
            elif element.name == 'h3':
                doc.add_heading(text_content, level=3)
            elif element.name == 'h4':
                doc.add_heading(text_content, level=4)
            
            # 处理普通段落
            elif element.name == 'p':
                add_safe_paragraph(text_content)
            
            # 处理列表项
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    li_text = li.get_text(separator='\n').strip()
                    if li_text:
                        add_safe_paragraph(li_text, prefix="• ")

    # 保存文件
    try:
        doc.save(docx_path)
        print(f"🎉 [大功告成] Word 标书已生成！")
        return True
    except Exception as e:
        print(f"❌ [错误] Word 文档保存失败，请检查文件是否被占用打开了！详细错误: {e}")
        return False

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    root_dir = os.path.dirname(current_dir)
    template_file = os.path.join(root_dir, "标书排版模板.docx") 

    if len(sys.argv) >= 3:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
        convert_md_to_docx(md_file, docx_file, template_file)
    else:
        print("💡 提示: 这个脚本是被主程序调用的。")
        print("单独测试请在终端输入: python src/md2word.py <你的md文件路径> <输出的docx路径>")
